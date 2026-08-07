"""Office exporters: DOCX and XLSX.

These are the formats the team actually pastes into client decks and briefs, so
they get real formatting rather than a text dump with a different extension.
"""

from __future__ import annotations

import io

from app.core.text import format_duration, format_timestamp
from app.services.export.base import ExportDocument, Exporter


class DocxExporter(Exporter):
    format = "docx"
    extension = "docx"
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    display_name = "Word document"
    combinable = True

    def _meta_table(self, doc, document: ExportDocument) -> None:  # noqa: ANN001
        meta_rows = [
            ("Creator", document.author or "—"),
            ("Platform", document.platform),
            ("Duration", format_duration(document.duration_seconds)),
            ("Language", document.language or "—"),
            ("Words", str(document.word_count)),
            ("Source", document.source_url),
        ]
        table = doc.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"
        for label, value in meta_rows:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = value
            cells[0].paragraphs[0].runs[0].bold = True

    def render(self, document: ExportDocument) -> bytes:
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        doc.add_heading(document.safe_title, level=1)
        self._meta_table(doc, document)

        doc.add_paragraph()
        doc.add_heading("Transcript", level=2)
        body = doc.add_paragraph(document.text or "(no speech detected)")
        body.style.font.size = Pt(11)

        if document.segments:
            doc.add_page_break()
            doc.add_heading("Timed segments", level=2)
            segment_table = doc.add_table(rows=1, cols=2)
            segment_table.style = "Light List Accent 1"
            header = segment_table.rows[0].cells
            header[0].text = "Time"
            header[1].text = "Text"
            for segment in document.segments:
                cells = segment_table.add_row().cells
                cells[0].text = format_timestamp(segment.start, separator=".")[:-4]
                cells[1].text = segment.text

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def render_many(self, documents: list[ExportDocument]) -> bytes:
        # One brief covering the whole batch, which is what gets pasted into a
        # deck. Each video starts on its own page so it can still be pulled out
        # individually. Segment tables are omitted for the reason given in the
        # Markdown exporter — export one transcript when the timings matter.
        from docx import Document

        doc = Document()
        doc.add_heading("Combined transcripts", level=0)
        doc.add_paragraph(f"{len(documents)} videos.")

        doc.add_heading("Contents", level=1)
        for number, document in enumerate(documents, start=1):
            doc.add_paragraph(f"{number}. {document.safe_title}", style="List Number")

        for document in documents:
            doc.add_page_break()
            doc.add_heading(document.safe_title, level=1)
            self._meta_table(doc, document)
            doc.add_paragraph()
            doc.add_paragraph(document.text or "(no speech detected)")

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()


class XlsxExporter(Exporter):
    format = "xlsx"
    extension = "xlsx"
    content_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    display_name = "Excel workbook"
    combinable = True

    _SUMMARY_HEADERS = (
        "Title", "Creator", "Platform", "Duration", "Language",
        "Words", "Provider", "Source URL", "Transcript",
    )
    _SEGMENT_HEADERS = ("Title", "#", "Start", "End", "Text")

    def _style_header(self, sheet, columns: int) -> None:  # noqa: ANN001
        from openpyxl.styles import Alignment, Font, PatternFill

        fill = PatternFill("solid", start_color="1F2937")
        for index in range(1, columns + 1):
            cell = sheet.cell(row=1, column=index)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = fill
            cell.alignment = Alignment(vertical="center")
        sheet.freeze_panes = "A2"

    def _build(self, documents: list[ExportDocument]) -> bytes:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment

        workbook = Workbook()

        summary = workbook.active
        summary.title = "Transcripts"
        summary.append(list(self._SUMMARY_HEADERS))
        for document in documents:
            summary.append(
                [
                    document.safe_title,
                    document.author or "",
                    document.platform,
                    format_duration(document.duration_seconds),
                    document.language or "",
                    document.word_count,
                    document.provider or "",
                    document.source_url,
                    document.text,
                ]
            )
        self._style_header(summary, len(self._SUMMARY_HEADERS))
        for column, width in zip("ABCDEFGHI", (46, 22, 12, 12, 10, 9, 14, 46, 90), strict=False):
            summary.column_dimensions[column].width = width
        for row in summary.iter_rows(min_row=2, min_col=9, max_col=9):
            for cell in row:
                cell.alignment = Alignment(wrap_text=True, vertical="top")

        segments_sheet = workbook.create_sheet("Segments")
        segments_sheet.append(list(self._SEGMENT_HEADERS))
        for document in documents:
            for segment in document.segments:
                segments_sheet.append(
                    [
                        document.safe_title,
                        segment.index + 1,
                        format_timestamp(segment.start, separator="."),
                        format_timestamp(segment.end, separator="."),
                        segment.text,
                    ]
                )
        self._style_header(segments_sheet, len(self._SEGMENT_HEADERS))
        for column, width in zip("ABCDE", (40, 6, 14, 14, 100), strict=False):
            segments_sheet.column_dimensions[column].width = width

        buffer = io.BytesIO()
        workbook.save(buffer)
        return buffer.getvalue()

    def render(self, document: ExportDocument) -> bytes:
        return self._build([document])

    def render_many(self, documents: list[ExportDocument]) -> bytes:
        # One workbook with every transcript beats a ZIP of workbooks — this is
        # the format people use precisely because they want to compare rows.
        return self._build(documents)
